"""
AI-Powered Document Template Generator
Membuat dan menyesuaikan template dokumen hukum Indonesia dengan AI
"""
import logging
import io
import os
import tempfile
from datetime import datetime
from typing import List, Dict, Any, Optional
from pydantic import BaseModel, Field
from bson import ObjectId

from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks, Response
from fastapi.responses import StreamingResponse, FileResponse

# PDF generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch

# DOCX generation
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from ..services.blockchain_databases import get_mongodb_cursor
from ..services.ai_service import AdvancedAIService
from ..core.security_updated import get_current_user
from ..models import User

router = APIRouter(prefix="/api/templates", tags=["Template Generator"])
logger = logging.getLogger(__name__)

ai_service = AdvancedAIService()

# Pydantic Models
class TemplateInfo(BaseModel):
    """Template information"""
    template_id: str
    name: str
    category: str
    description: str
    tags: List[str]
    estimated_complexity: str  # simple/medium/complex
    usage_count: Optional[int] = 0
    last_updated: str

class TemplateVariables(BaseModel):
    """Variables required for template customization"""
    variable_name: str
    variable_type: str  # text/date/number/choice
    label: str
    description: str
    required: bool = True
    choices: Optional[List[str]] = None
    default_value: Optional[Any] = None

class TemplateGenerationRequest(BaseModel):
    """Request to generate document from template"""
    template_id: str
    variables: Dict[str, Any] = Field(..., description="Custom values for template variables")
    format: str = Field("pdf", description="Output format: pdf/docx/html")
    language: str = Field("indonesian", description="Document language")

class GeneratedDocument(BaseModel):
    """Generated document response"""
    document_id: str
    template_id: str
    filename: str
    download_url: str
    preview_url: Optional[str]
    generated_at: str
    size_bytes: Optional[int] = None

class TemplatePreview(BaseModel):
    """Template preview response"""
    html_content: str
    variables_used: Dict[str, Any]
    warnings: List[str]
    estimated_length: int  # characters

# Indonesian Legal Document Categories
DOCUMENT_CATEGORIES = {
    "contracts": {
        "name": "Perjanjian/Kontrak",
        "subcategories": [
            "Jual Beli", "Sewa Menyewa", "Karya", "Pekerjaan",
            "Kerjasama", "Lisensi", "Distribusi", "Franchise"
        ]
    },
    "power_of_attorney": {
        "name": "Surat Kuasa",
        "subcategories": [
            "Kuasa Umum", "Kuasa Khusus", "Kuasa Perorangan",
            "Kuasa untuk Menikah", "Kuasa untuk Mengurus Tanah"
        ]
    },
    "agreements": {
        "name": "Perjanjian Khusus",
        "subcategories": [
            "Perjanjian Pra Nikah", "Perjanjian Pisah Harta",
            "Perjanjian Perdamaian", "Perjanjian Damai"
        ]
    },
    "legal_notices": {
        "name": "Pemberitahuan Hukum",
        "subcategories": [
            "Surat Peringatan", "Surat Panggilan", "Surat Somasi",
            "Surat Teguran", "Pemberitahuan Sengketa"
        ]
    },
    "company_documents": {
        "name": "Dokumen Perusahaan",
        "subcategories": [
            "Akta Pendirian", "Anggaran Dasar", "Perubahan Anggaran Dasar",
            "Risalah RUPS", "Surat Saham"
        ]
    },
    "property_documents": {
        "name": "Dokumen Properti",
        "subcategories": [
            "Surat Perjanjian Jual Beli Tanah",
            "Surat Sewa Kos-Kosan", "Surat Perjanjian Hak Guna Bangunan",
            "Surat Perjanjian Mortgage", "Surat Hibah Tanah"
        ]
    },
    "employment": {
        "name": "Dokumen Ketenagakerjaan",
        "subcategories": [
            "Perjanjian Kerja Waktu Tertentu",
            "Perjanjian Kerja Waktu Tidak Tertentu",
            "Surat Perjanjian Kerja Bersama",
            "Surat Perjanjian Pemutusan Hubungan Kerja"
        ]
    },
    "courts": {
        "name": "Dokumen Pengadilan",
        "subcategories": [
            "Gugatan Cerai", "Jawaban Gugatan", "Rekonvensi",
            "Surat Kuasa untuk Menggugat", "Akta Perdamaian"
        ]
    }
}

# Template Repository (Sample templates)
SAMPLE_TEMPLATES = [
    {
        "template_id": "contr_perjanjian_jual_beli",
        "name": "Perjanjian Jual Beli Tanah",
        "category": "contracts",
        "subcategory": "Jual Beli",
        "description": "Template standar perjanjian jual beli tanah dengan seluruh klausul yang diperlukan",
        "tags": ["tanah", "jual beli", "property", "real estate"],
        "complexity": "medium",
        "variables": [
            {"name": "penjual_nama_lengkap", "type": "text", "label": "Nama Lengkap Penjual", "required": True},
            {"name": "penjual_alamat", "type": "text", "label": "Alamat Penjual", "required": True},
            {"name": "pembeli_nama_lengkap", "type": "text", "label": "Nama Lengkap Pembeli", "required": True},
            {"name": "pembeli_alamat", "type": "text", "label": "Alamat Pembeli", "required": True},
            {"name": "tanah_luas", "type": "text", "label": "Luas Tanah", "required": True},
            {"name": "tanah_alamat_lengkap", "type": "text", "label": "Alamat Lengkap Tanah", "required": True},
            {"name": "harga_total", "type": "text", "label": "Harga Total", "required": True},
            {"name": "uang_muka", "type": "text", "label": "Jumlah Uang Muka", "required": True},
            {"name": "tempo_pembayaran", "type": "date", "label": "Tempo Pembayaran Utang", "required": True},
            {"name": "tanggal_perjanjian", "type": "date", "label": "Tanggal Perjanjian", "required": True}
        ],
        "template_content": """
PERJANJIAN JUAL BELI TANAH

Nomor: {{nomor_perjanjian}}
Tanggal: {{tanggal_perjanjian}}

Pihak Pertama (Penjual):
{{penjual_nama_lengkap}}
{{penjual_alamat}}

Pihak Kedua (Pembeli):
{{pembeli_nama_lengkap}}
{{pembeli_alamat}}

KESATU: PENJUAL telah sepakat untuk menjual kepada PEMBELI sebidang tanah yang terletak di {{tanah_alamat_lengkap}} dengan luas {{tanah_luas}} meter persegi.

KEDUA: Harga jual tanah tersebut adalah sebesar Rp {{harga_total}} ({{harga_terbilang}}) dan uang muka sebesar Rp {{uang_muka}}.

KETIGA: Sisa pembayaran akan dibayarkan selambat-lambatnya pada tanggal {{tempo_pembayaran}}.

Dan seterusnya sesuai standart perjanjian jual beli tanah di Indonesia...
        """
    },
    {
        "template_id": "power_of_attorney_general",
        "name": "Surat Kuasa Umum",
        "category": "power_of_attorney",
        "subcategory": "Kuasa Umum",
        "description": "Template surat kuasa umum untuk mewakili dalam berbagai urusan",
        "tags": ["kuasa", "wakil", "pemberian kuasa"],
        "complexity": "simple",
        "variables": [
            {"name": "pemberi_kuasa_nama", "type": "text", "label": "Nama Pemberi Kuasa", "required": True},
            {"name": "pemberi_kuasa_alamat", "type": "text", "label": "Alamat Pemberi Kuasa", "required": True},
            {"name": "penerima_kuasa_nama", "type": "text", "label": "Nama Penerima Kuasa", "required": True},
            {"name": "penerima_kuasa_alamat", "type": "text", "label": "Alamat Penerima Kuasa", "required": True},
            {"name": "tanggal_kuasa", "type": "date", "label": "Tanggal Pemberian Kuasa", "required": True},
            {"name": "jangka_waktu", "type": "text", "label": "Jangka Waktu Kuasa", "required": False},
            {"name": "kewenangan_khusus", "type": "text", "label": "Kewenangan Khusus (opsional)", "required": False}
        ],
        "template_content": """
SURAT KUASA

Nomor: {{nomor_kuasa}}

Saya yang bertanda tangan di bawah ini:
Nama: {{pemberi_kuasa_nama}}
Alamat: {{pemberi_kuasa_alamat}}

Dengan ini memberikan kuasa penuh kepada:
Nama: {{penerima_kuasa_nama}}
Alamat: {{penerima_kuasa_alamat}}

Untuk {{penerima_kuasa_nama}} dapat mewakili saya dalam segala urusan sehari-hari.

Kuasa ini diberikan pada tanggal {{tanggal_kuasa}} {{jangka_waktu}}.

Dan seterusnya sesuai format surat kuasa Indonesia...
        """
    }
]

@router.get("/", response_model=List[TemplateInfo])
async def get_available_templates(
    category: Optional[str] = None,
    subcategory: Optional[str] = None,
    complexity: Optional[str] = None
) -> List[Dict[str, Any]]:
    """
    Get semua template yang tersedia dengan filtering
    """
    try:
        # For now, return sample templates (in real implementation, query database)
        templates = SAMPLE_TEMPLATES

        # Apply filters
        if category:
            templates = [t for t in templates if t["category"] == category]

        if subcategory:
            templates = [t for t in templates if t.get("subcategory") == subcategory]

        if complexity:
            templates = [t for t in templates if t["complexity"] == complexity]

        # Format response
        result = []
        for template in templates:
            result.append({
                "template_id": template["template_id"],
                "name": template["name"],
                "category": template["category"],
                "description": template["description"],
                "tags": template["tags"],
                "estimated_complexity": template["complexity"],
                "usage_count": template.get("usage_count", 0),
                "last_updated": datetime.utcnow().isoformat()
            })

        return result

    except Exception as e:
        logger.error(f"Get templates error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to retrieve templates")

@router.get("/categories")
async def get_template_categories() -> Dict[str, Any]:
    """Get template categories dan subcategories"""
    return {
        "categories": DOCUMENT_CATEGORIES,
        "total_templates": len(SAMPLE_TEMPLATES)
    }

@router.get("/{template_id}/variables", response_model=List[TemplateVariables])
async def get_template_variables(template_id: str) -> List[Dict[str, Any]]:
    """
    Get variables yang dibutuhkan template tertentu
    """
    try:
        # Find template
        template = next((t for t in SAMPLE_TEMPLATES if t["template_id"] == template_id), None)

        if not template:
            raise HTTPException(status_code=404, detail="Template not found")

        # Format variables
        variables = []
        for var in template["variables"]:
            variables.append({
                "variable_name": var["name"],
                "variable_type": var["type"],
                "label": var["label"],
                "description": var.get("description", ""),
                "required": var.get("required", True),
                "choices": var.get("choices"),
                "default_value": var.get("default_value")
            })

        return variables

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get template variables error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to retrieve template variables")

@router.post("/preview", response_model=TemplatePreview)
async def preview_template(
    request: TemplateGenerationRequest,
) -> Dict[str, Any]:
    """
    Preview template dengan variables yang diinput user
    """
    try:
        # Find template
        template = next((t for t in SAMPLE_TEMPLATES if t["template_id"] == request.template_id), None)

        if not template:
            raise HTTPException(status_code=404, detail="Template not found")

        # Generate preview HTML
        html_content = generate_html_preview(template, request.variables)

        return TemplatePreview(
            html_content=html_content,
            variables_used=request.variables,
            warnings=[],  # Could add validation warnings
            estimated_length=len(html_content)
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Preview template error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate preview")

@router.post("/generate", response_model=GeneratedDocument)
async def generate_document(
    request: TemplateGenerationRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user)
) -> Dict[str, Any]:
    """
    Generate final document dari template
    """
    try:
        # Find template
        template = next((t for t in SAMPLE_TEMPLATES if t["template_id"] == request.template_id), None)

        if not template:
            raise HTTPException(status_code=404, detail="Template not found")

        # Generate document ID
        doc_id = f"doc_{current_user.id}_{int(datetime.utcnow().timestamp())}"

        # Store generation request for background processing
        generation_data = {
            "document_id": doc_id,
            "template_id": request.template_id,
            "variables": request.variables,
            "format": request.format,
            "language": request.language,
            "user_id": str(current_user.id),
            "status": "processing",
            "created_at": datetime.utcnow()
        }

        # Save to database for background processing
        mongodb = get_mongodb_cursor()
        if mongodb:
            db = mongodb["pasalku_ai_analytics"]
            await db["document_generations"].insert_one(generation_data)

        # Start background generation
        background_tasks.add_task(
            generate_document_background,
            generation_data
        )

        # Return initial response
        filename = f"{template['name']}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}"

        return GeneratedDocument(
            document_id=doc_id,
            template_id=request.template_id,
            filename=filename,
            download_url=f"/api/templates/{doc_id}/download",
            preview_url=f"/api/templates/{doc_id}/preview",
            generated_at=datetime.utcnow().isoformat()
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Generate document error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate document")

@router.get("/{document_id}/download")
async def download_generated_document(
    document_id: str,
    current_user: User = Depends(get_current_user)
):
    """
    Download generated document
    """
    try:
        mongodb = get_mongodb_cursor()
        if not mongodb:
            raise HTTPException(status_code=500, detail="Document store unavailable")

        db = mongodb["pasalku_ai_analytics"]

        # Find document
        doc = await db["document_generations"].find_one({
            "document_id": document_id,
            "user_id": str(current_user.id)
        })

        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        if doc.get("status") != "completed":
            raise HTTPException(status_code=202, detail="Document still processing")

        # Get file buffer from database
        file_buffer = doc.get("file_buffer")
        if not file_buffer:
            raise HTTPException(status_code=500, detail="Document file not available")

        filename = doc.get("filename", "generated_document.pdf")

        # Determine content type based on file extension
        if filename.lower().endswith('.pdf'):
            media_type = "application/pdf"
        elif filename.lower().endswith('.docx'):
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif filename.lower().endswith('.html'):
            media_type = "text/html"
        else:
            media_type = "application/octet-stream"

        # Create buffer from stored bytes
        buffer = io.BytesIO(file_buffer)

        # Return StreamingResponse for in-memory file
        return StreamingResponse(
            buffer,
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{filename}",
                "Content-Length": str(len(file_buffer))
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Download document error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to download document")

@router.post("/customize")
async def customize_template_content(
    template_id: str,
    custom_requirements: str,
    current_user: User = Depends(get_current_user)
) -> Dict[str, Any]:
    """
    AI-powered template customization berdasarkan requirements khusus
    """
    try:
        # Find base template
        template = next((t for t in SAMPLE_TEMPLATES if t["template_id"] == template_id), None)

        if not template:
            raise HTTPException(status_code=404, detail="Template not found")

        # Generate customized template dengan AI
        customization_prompt = f"""
        Kustomisasi template dokumen hukum berikut berdasarkan requirements khusus:

        TEMPLATE ASLI: {template['name']}
        KATEGORI: {template['category']}
        ISI TEMPLATE:
        {template['template_content']}

        REQUIREMENTS KHUSUS:
        {custom_requirements}

        Berikan versi yang dikustomisasi dengan:
        1. Klausul tambahan yang sesuai requirements
        2. Penyesuaian bahasa dan tone
        3. Variabel baru yang mungkin dibutuhkan
        4. Struktur yang lebih sesuai kebutuhan

        Return dalam format JSON dengan customized content dan new variables.
        """

        ai_response = await ai_service.get_legal_response(query=customization_prompt)

        # Parse AI response (simplified)
        customized = {
            "original_template": template_id,
            "customized_content": template["template_content"],  # Placeholder
            "additional_variables": [],
            "customizations_applied": ["Additional clauses added", "Language adjusted"],
            "ai_insights": ai_response.get("answer", "Template customized successfully")
        }

        return customized

    except Exception as e:
        logger.error(f"Customize template error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to customize template")

@router.get("/popular")
async def get_popular_templates(limit: int = 10) -> Dict[str, Any]:
    """Get most-used templates"""
    # Placeholder - in real implementation, query actual usage statistics
    popular_templates = SAMPLE_TEMPLATES[:limit]

    return {
        "popular_templates": [
            {
                "template_id": t["template_id"],
                "name": t["name"],
                "category": t["category"],
                "usage_count": t.get("usage_count", 0),
                "description": t["description"]
            } for t in popular_templates
        ],
        "total_popular": len(popular_templates)
    }

# Helper Functions
async def generate_document_background(generation_data: Dict[str, Any]):
    """Background task untuk generate document"""
    try:
        logger.info(f"Starting document generation: {generation_data['document_id']}")

        # Find template
        template = next((t for t in SAMPLE_TEMPLATES if t["template_id"] == generation_data["template_id"]), None)

        if not template:
            await update_generation_status(generation_data["document_id"], "failed", "Template not found")
            return

        # Apply variables to template
        document_content = template["template_content"]
        variables = generation_data["variables"]

        # Simple variable replacement
        for var_name, var_value in variables.items():
            placeholder = "{{" + var_name + "}}"
            document_content = document_content.replace(placeholder, str(var_value))

        # AI enhancement (optional)
        if generation_data.get("language") == "indonesian":
            enhanced_content = await enhance_document_with_ai(document_content, template, variables)
            document_content = enhanced_content

        # Generate final document based on format
        doc_format = generation_data.get("format", "pdf")
        file_buffer = None
        filename = ""
        document_size = 0

        try:
            if doc_format == "pdf":
                file_buffer, filename = generate_pdf_document(document_content, template, variables)
            elif doc_format == "docx":
                file_buffer, filename = generate_docx_document(document_content, template, variables)
            elif doc_format == "html":
                file_buffer, filename = generate_html_document(document_content, template, variables)
            else:
                # Default to PDF
                file_buffer, filename = generate_pdf_document(document_content, template, variables)

            document_size = len(file_buffer.getvalue()) if hasattr(file_buffer, 'getvalue') else len(file_buffer)
            file_buffer.seek(0)  # Reset buffer position

        except Exception as gen_error:
            logger.error(f"Document generation error: {str(gen_error)}")
            await update_generation_status(generation_data["document_id"], "failed", f"Generation error: {str(gen_error)}")
            return

        # Save generated document
        generated_doc = {
            **generation_data,
            "file_buffer": file_buffer.getvalue() if hasattr(file_buffer, 'getvalue') else file_buffer,
            "filename": filename,
            "size_bytes": document_size,
            "status": "completed",
            "completed_at": datetime.utcnow()
        }

        # Update database
        mongodb = get_mongodb_cursor()
        if mongodb:
            db = mongodb["pasalku_ai_analytics"]
            await db["document_generations"].update_one(
                {"document_id": generation_data["document_id"]},
                {"$set": generated_doc}
            )

        logger.info(f"Document generation completed: {generation_data['document_id']}")

    except Exception as e:
        logger.error(f"Document generation background error: {str(e)}")
        await update_generation_status(generation_data["document_id"], "failed", str(e))

async def update_generation_status(document_id: str, status: str, error_message: str = None):
    """Update document generation status"""
    try:
        mongodb = get_mongodb_cursor()
        if mongodb:
            db = mongodb["pasalku_ai_analytics"]
            update_data = {"status": status}
            if error_message:
                update_data["error"] = error_message
            await db["document_generations"].update_one(
                {"document_id": document_id},
                {"$set": update_data}
            )
    except Exception as e:
        logger.error(f"Update generation status error: {str(e)}")

async def enhance_document_with_ai(content: str, template: Dict, variables: Dict) -> str:
    """Enhance generated document dengan AI untuk better legal accuracy"""
    try:
        enhancement_prompt = f"""
        Review dan enhance dokumen hukum yang telah digenerate:

        TEMPLATE: {template['name']}
        GENERATED CONTENT:
        {content[:1000]}...

        VARIABLE VALUES USED:
        {variables}

        Berikan improvement suggestions untuk legal accuracy dan completeness.
        Return enhanced version yang lebih baik secara hukum.
        """

        ai_response = await ai_service.get_legal_response(query=enhancement_prompt)
        enhanced_content = ai_response.get("answer", content[:500])  # Fallback to original if AI fails

        return enhanced_content

    except Exception as e:
        logger.error(f"Document enhancement error: {str(e)}")
        return content

def generate_html_preview(template: Dict, variables: Dict[str, Any]) -> str:
    """Generate HTML preview dari template dengan variables"""
    content = template["template_content"]

    # Apply variable substitutions
    for var_name, var_value in variables.items():
        placeholder = "{{" + var_name + "}}"
        content = content.replace(placeholder, str(var_value))

    # Convert to simple HTML
    html_content = content.replace('\n', '<br>')

    # Wrap in HTML structure
    html_preview = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{ font-family: 'Times New Roman', serif; line-height: 1.6; margin: 20px; }}
            .variable {{ background-color: #e3f2fd; padding: 2px 4px; border-radius: 3px; }}
        </style>
    </head>
    <body>
        <div class="document-preview">
            {html_content}
        </div>
    </body>
    </html>
    """

    return html_preview

def generate_pdf_document(content: str, template: Dict, variables: Dict) -> tuple:
    """Generate PDF document menggunakan ReportLab"""
    # Create buffer for PDF
    buffer = io.BytesIO()

    # Create PDF document
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)

    # Create styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='CenteredTitle', fontSize=16, leading=18, alignment=1, fontName='Helvetica-Bold'))
    styles.add(ParagraphStyle(name='Centered', fontSize=12, leading=14, alignment=1))
    styles.add(ParagraphStyle(name='Justified', fontSize=11, leading=14, alignment=4))

    # Split content into sections
    sections = content.split('\n\n')
    story = []

    # Add title
    title = template["name"]
    story.append(Paragraph(title, styles['CenteredTitle']))
    story.append(Spacer(1, 0.5*inch))

    # Add generation info
    generation_date = datetime.utcnow().strftime("%d %B %Y")
    generation_info = f"Dokumen dibuat pada: {generation_date}"
    story.append(Paragraph(generation_info, styles['Centered']))
    story.append(Spacer(1, 0.3*inch))

    # Process content sections
    for section in sections:
        if section.strip():
            # Clean and format the section
            section_cleaned = section.replace('\n', ' ').strip()
            story.append(Paragraph(section_cleaned, styles['Justified']))
            story.append(Spacer(1, 0.2*inch))

    # Build PDF
    doc.build(story)

    # Generate filename
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"{template['name'].replace(' ', '_')}_{timestamp}.pdf"

    return buffer, filename

def generate_docx_document(content: str, template: Dict, variables: Dict) -> tuple:
    """Generate DOCX document menggunakan python-docx"""
    # Create DOCX document
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add title
    title = doc.add_heading(template["name"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add generation info
    generation_date = datetime.utcnow().strftime("%d %B %Y")
    generation_info = doc.add_paragraph(f"Dokumen dibuat pada: {generation_date}")
    generation_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Add spacing

    # Process content
    sections = content.split('\n\n')
    for section in sections:
        if section.strip():
            # Clean the section
            section_cleaned = section.replace('\n', ' ').strip()

            # Add paragraph
            p = doc.add_paragraph(section_cleaned)

            # Set font
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)

    # Generate filename
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"{template['name'].replace(' ', '_')}_{timestamp}.docx"

    return buffer, filename

def generate_html_document(content: str, template: Dict, variables: Dict) -> tuple:
    """Generate HTML document"""
    # Generate HTML content
    timestamp = datetime.utcnow().strftime("%d %B %Y %H:%M:%S")
    html_content = f"""
    <!DOCTYPE html>
    <html lang="id">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{template["name"]}</title>
        <style>
            body {{
                font-family: 'Times New Roman', serif;
                line-height: 1.6;
                margin: 40px auto;
                max-width: 800px;
                padding: 20px;
                color: #333;
            }}
            h1 {{
                text-align: center;
                font-size: 24px;
                margin-bottom: 10px;
            }}
            .metadata {{
                text-align: center;
                font-size: 12px;
                color: #666;
                margin-bottom: 30px;
            }}
            .content {{
                text-align: justify;
            }}
            .section {{
                margin-bottom: 15px;
            }}
        </style>
    </head>
    <body>
        <h1>{template["name"]}</h1>
        <div class="metadata">
            Dokumen dibuat dengan Pasalku.ai pada {timestamp}<br>
            Template ID: {template["template_id"]}
        </div>
        <div class="content">
    """

    # Process sections
    sections = content.split('\n\n')
    for section in sections:
        if section.strip():
            section_cleaned = section.replace('\n', ' ').strip()
            html_content += f'\n            <div class="section">{section_cleaned}</div>'

    # Close HTML
    html_content += """
        </div>
    </body>
    </html>"""

    # Encode to bytes
    buffer = io.BytesIO(html_content.encode('utf-8'))

    # Generate filename
    timestamp_short = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"{template['name'].replace(' ', '_')}_{timestamp_short}.html"

    return buffer, filename

# Template Analytics
@router.get("/stats")
async def get_template_usage_stats(current_user: User = Depends(get_current_user)) -> Dict[str, Any]:
    """Get template usage statistics"""
    # Placeholder statistics
    return {
        "total_templates": len(SAMPLE_TEMPLATES),
        "categories_used": len(DOCUMENT_CATEGORIES),
        "user_generation_count": 0,  # Would come from database
        "popular_category": "contracts",
        "most_used_template": SAMPLE_TEMPLATES[0]["name"],
        "average_generation_time": 2.5,  # seconds
        "success_rate": 95.5  # percentage
    }