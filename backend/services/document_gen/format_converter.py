"""
Format Converter - Convert dokumen ke berbagai format

Mengkonversi dokumen yang sudah diisi ke:
- DOCX (Microsoft Word)
- PDF
- HTML
- Markdown
"""

from dataclasses import dataclass
from typing import Optional, List
from pathlib import Path
import tempfile
from datetime import datetime
from enum import Enum


class DocumentFormat(Enum):
    """Format dokumen yang didukung"""
    DOCX = "docx"
    PDF = "pdf"
    HTML = "html"
    MARKDOWN = "markdown"
    TXT = "txt"


@dataclass
class ConversionResult:
    """Result of format conversion"""
    success: bool
    format: DocumentFormat
    file_path: Optional[str] = None
    file_content: Optional[bytes] = None
    error: Optional[str] = None


class FormatConverter:
    """
    Convert documents ke berbagai format
    """
    
    def __init__(self):
        self.temp_dir = Path(tempfile.gettempdir()) / "pasalku_docs"
        self.temp_dir.mkdir(exist_ok=True)
        
        # Check available libraries
        self.docx_available = False
        self.pdf_available = False
        
        try:
            import docx
            self.docx_available = True
        except ImportError:
            pass
        
        try:
            from reportlab.lib.pagesizes import A4
            self.pdf_available = True
        except ImportError:
            pass
    
    def convert(
        self,
        content: str,
        target_format: DocumentFormat,
        title: Optional[str] = None,
        metadata: Optional[dict] = None
    ) -> ConversionResult:
        """
        Convert content ke format yang diminta
        
        Args:
            content: Document content (plain text)
            target_format: Target format
            title: Document title
            metadata: Additional metadata
        
        Returns:
            ConversionResult
        """
        try:
            if target_format == DocumentFormat.DOCX:
                return self._convert_to_docx(content, title, metadata)
            elif target_format == DocumentFormat.PDF:
                return self._convert_to_pdf(content, title, metadata)
            elif target_format == DocumentFormat.HTML:
                return self._convert_to_html(content, title, metadata)
            elif target_format == DocumentFormat.MARKDOWN:
                return self._convert_to_markdown(content, title, metadata)
            elif target_format == DocumentFormat.TXT:
                return self._convert_to_txt(content)
            else:
                return ConversionResult(
                    success=False,
                    format=target_format,
                    error=f"Unsupported format: {target_format}"
                )
        except Exception as e:
            return ConversionResult(
                success=False,
                format=target_format,
                error=str(e)
            )
    
    def _convert_to_docx(self, content: str, title: Optional[str], metadata: Optional[dict]) -> ConversionResult:
        """Convert to DOCX format"""
        if not self.docx_available:
            return ConversionResult(
                success=False,
                format=DocumentFormat.DOCX,
                error="python-docx library not available. Install with: pip install python-docx"
            )
        
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # Add title if provided
            if title:
                heading = doc.add_heading(title, level=0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            if metadata:
                self._add_docx_metadata(doc, metadata)
            
            # Add content
            # Split by paragraphs
            paragraphs = content.split('\n\n')
            
            for para_text in paragraphs:
                if not para_text.strip():
                    continue
                
                # Check if it's a heading (all caps or starts with "Pasal")
                if para_text.strip().isupper() or para_text.strip().startswith("Pasal"):
                    p = doc.add_heading(para_text.strip(), level=1)
                else:
                    p = doc.add_paragraph(para_text.strip())
                
                # Set font
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            # Save to temp file
            filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            file_path = self.temp_dir / filename
            doc.save(str(file_path))
            
            # Read file content
            with open(file_path, 'rb') as f:
                file_content = f.read()
            
            return ConversionResult(
                success=True,
                format=DocumentFormat.DOCX,
                file_path=str(file_path),
                file_content=file_content
            )
        
        except Exception as e:
            return ConversionResult(
                success=False,
                format=DocumentFormat.DOCX,
                error=f"Error creating DOCX: {str(e)}"
            )
    
    def _add_docx_metadata(self, doc, metadata: dict):
        """Add metadata to DOCX document"""
        core_properties = doc.core_properties
        
        if "author" in metadata:
            core_properties.author = metadata["author"]
        if "subject" in metadata:
            core_properties.subject = metadata["subject"]
        if "keywords" in metadata:
            core_properties.keywords = metadata["keywords"]
        if "comments" in metadata:
            core_properties.comments = metadata["comments"]
        
        core_properties.created = datetime.now()
        core_properties.modified = datetime.now()
    
    def _convert_to_pdf(self, content: str, title: Optional[str], metadata: Optional[dict]) -> ConversionResult:
        """Convert to PDF format"""
        if not self.pdf_available:
            return ConversionResult(
                success=False,
                format=DocumentFormat.PDF,
                error="reportlab library not available. Install with: pip install reportlab"
            )
        
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # Create PDF
            filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = self.temp_dir / filename
            
            doc = SimpleDocTemplate(
                str(file_path),
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Build story
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                alignment=TA_CENTER,
                spaceAfter=30
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=12,
                alignment=TA_JUSTIFY,
                spaceAfter=12
            )
            
            # Add title
            if title:
                story.append(Paragraph(title, title_style))
                story.append(Spacer(1, 0.2 * inch))
            
            # Add content
            paragraphs = content.split('\n\n')
            
            for para_text in paragraphs:
                if not para_text.strip():
                    continue
                
                # Check if heading
                if para_text.strip().isupper() or para_text.strip().startswith("Pasal"):
                    p = Paragraph(para_text.strip(), styles['Heading2'])
                else:
                    p = Paragraph(para_text.strip(), body_style)
                
                story.append(p)
            
            # Build PDF
            doc.build(story)
            
            # Read file content
            with open(file_path, 'rb') as f:
                file_content = f.read()
            
            return ConversionResult(
                success=True,
                format=DocumentFormat.PDF,
                file_path=str(file_path),
                file_content=file_content
            )
        
        except Exception as e:
            return ConversionResult(
                success=False,
                format=DocumentFormat.PDF,
                error=f"Error creating PDF: {str(e)}"
            )
    
    def _convert_to_html(self, content: str, title: Optional[str], metadata: Optional[dict]) -> ConversionResult:
        """Convert to HTML format"""
        try:
            # Build HTML
            html_parts = []
            
            # HTML header
            html_parts.append("<!DOCTYPE html>")
            html_parts.append("<html lang='id'>")
            html_parts.append("<head>")
            html_parts.append("    <meta charset='UTF-8'>")
            html_parts.append("    <meta name='viewport' content='width=device-width, initial-scale=1.0'>")
            
            if title:
                html_parts.append(f"    <title>{title}</title>")
            
            # CSS
            html_parts.append("    <style>")
            html_parts.append("        body {")
            html_parts.append("            font-family: 'Times New Roman', serif;")
            html_parts.append("            font-size: 12pt;")
            html_parts.append("            line-height: 1.6;")
            html_parts.append("            max-width: 8.5in;")
            html_parts.append("            margin: 1in auto;")
            html_parts.append("            padding: 20px;")
            html_parts.append("            background: white;")
            html_parts.append("        }")
            html_parts.append("        h1 {")
            html_parts.append("            text-align: center;")
            html_parts.append("            font-size: 16pt;")
            html_parts.append("            margin-bottom: 30px;")
            html_parts.append("        }")
            html_parts.append("        h2 {")
            html_parts.append("            font-size: 14pt;")
            html_parts.append("            margin-top: 20px;")
            html_parts.append("        }")
            html_parts.append("        p {")
            html_parts.append("            text-align: justify;")
            html_parts.append("            margin-bottom: 12px;")
            html_parts.append("        }")
            html_parts.append("        @media print {")
            html_parts.append("            body { margin: 1in; }")
            html_parts.append("        }")
            html_parts.append("    </style>")
            html_parts.append("</head>")
            html_parts.append("<body>")
            
            # Title
            if title:
                html_parts.append(f"    <h1>{title}</h1>")
            
            # Content
            paragraphs = content.split('\n\n')
            
            for para_text in paragraphs:
                if not para_text.strip():
                    continue
                
                # Escape HTML
                para_text = para_text.replace('<', '&lt;').replace('>', '&gt;')
                
                # Check if heading
                if para_text.strip().isupper() or para_text.strip().startswith("Pasal"):
                    html_parts.append(f"    <h2>{para_text.strip()}</h2>")
                else:
                    html_parts.append(f"    <p>{para_text.strip()}</p>")
            
            html_parts.append("</body>")
            html_parts.append("</html>")
            
            html_content = "\n".join(html_parts)
            
            # Save to file
            filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            file_path = self.temp_dir / filename
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return ConversionResult(
                success=True,
                format=DocumentFormat.HTML,
                file_path=str(file_path),
                file_content=html_content.encode('utf-8')
            )
        
        except Exception as e:
            return ConversionResult(
                success=False,
                format=DocumentFormat.HTML,
                error=f"Error creating HTML: {str(e)}"
            )
    
    def _convert_to_markdown(self, content: str, title: Optional[str], metadata: Optional[dict]) -> ConversionResult:
        """Convert to Markdown format"""
        try:
            md_parts = []
            
            # Title
            if title:
                md_parts.append(f"# {title}")
                md_parts.append("")
            
            # Metadata
            if metadata:
                md_parts.append("---")
                for key, value in metadata.items():
                    md_parts.append(f"{key}: {value}")
                md_parts.append("---")
                md_parts.append("")
            
            # Content
            paragraphs = content.split('\n\n')
            
            for para_text in paragraphs:
                if not para_text.strip():
                    continue
                
                # Check if heading
                if para_text.strip().isupper() or para_text.strip().startswith("Pasal"):
                    md_parts.append(f"## {para_text.strip()}")
                else:
                    md_parts.append(para_text.strip())
                
                md_parts.append("")
            
            md_content = "\n".join(md_parts)
            
            # Save to file
            filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
            file_path = self.temp_dir / filename
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            return ConversionResult(
                success=True,
                format=DocumentFormat.MARKDOWN,
                file_path=str(file_path),
                file_content=md_content.encode('utf-8')
            )
        
        except Exception as e:
            return ConversionResult(
                success=False,
                format=DocumentFormat.MARKDOWN,
                error=f"Error creating Markdown: {str(e)}"
            )
    
    def _convert_to_txt(self, content: str) -> ConversionResult:
        """Convert to plain text"""
        try:
            filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            file_path = self.temp_dir / filename
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return ConversionResult(
                success=True,
                format=DocumentFormat.TXT,
                file_path=str(file_path),
                file_content=content.encode('utf-8')
            )
        
        except Exception as e:
            return ConversionResult(
                success=False,
                format=DocumentFormat.TXT,
                error=f"Error creating TXT: {str(e)}"
            )
